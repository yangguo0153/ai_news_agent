"""
Word 报销材料整合器
"""

import os
from pathlib import Path
from typing import List, Tuple
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    from PIL import Image
    HAS_PIL = True
except ImportError:
    HAS_PIL = False


class WordGenerator:
    """Word 报销材料整合器"""
    
    def __init__(self):
        if not HAS_DOCX:
            raise ImportError("请安装 python-docx: pip install python-docx")
        
        self.doc = Document()
        self._setup_document()
    
    def _setup_document(self):
        """设置文档基本格式"""
        # 设置页边距
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)
    
    def add_title(self, title: str):
        """添加标题"""
        heading = self.doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def add_section_title(self, title: str):
        """添加章节标题"""
        self.doc.add_heading(title, level=1)
    
    def add_subsection_title(self, title: str):
        """添加子标题"""
        self.doc.add_heading(title, level=2)
    
    def add_images_to_page(self, image_paths: List[str], title: str = "", max_per_row: int = 2):
        """将多张图片添加到文档（一页最多放2张）"""
        if title:
            self.add_subsection_title(title)
        
        # 计算每张图片的宽度
        page_width = 16  # cm, 减去边距
        img_width = page_width / max_per_row - 0.5  # 留一点间距
        
        for i in range(0, len(image_paths), max_per_row):
            batch = image_paths[i:i + max_per_row]
            
            # 创建表格来放置图片（无边框）
            if len(batch) > 1:
                table = self.doc.add_table(rows=1, cols=len(batch))
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                for col_idx, img_path in enumerate(batch):
                    cell = table.cell(0, col_idx)
                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run()
                    try:
                        run.add_picture(img_path, width=Cm(img_width))
                    except Exception as e:
                        paragraph.add_run(f"[图片加载失败: {Path(img_path).name}]")
            else:
                # 单张图片居中显示
                paragraph = self.doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                try:
                    run.add_picture(batch[0], width=Cm(12))
                except Exception as e:
                    paragraph.add_run(f"[图片加载失败: {Path(batch[0]).name}]")
            
            self.doc.add_paragraph()  # 添加空行
    
    def add_pdf_screenshots(self, pdf_paths: List[str], title: str = "", max_per_row: int = 2):
        """将 PDF 转为图片后添加到文档"""
        if not HAS_PYMUPDF:
            self.doc.add_paragraph(f"⚠️ 无法添加 PDF (PyMuPDF 未安装)")
            return

        temp_images = []

        for pdf_path in pdf_paths:
            try:
                doc = fitz.open(pdf_path)
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    # 渲染为图片
                    mat = fitz.Matrix(2, 2)  # 放大2倍以提高清晰度
                    pix = page.get_pixmap(matrix=mat)

                    # 保存为临时图片
                    temp_img_path = f"/tmp/pdf_page_{Path(pdf_path).stem}_{page_num}.png"
                    pix.save(temp_img_path)
                    temp_images.append(temp_img_path)
                doc.close()
            except Exception as e:
                print(f"⚠️  PDF 转换失败 {pdf_path}: {e}")

        # 添加图片到文档
        if temp_images:
            self.add_images_to_page(temp_images, "", max_per_row)

        # 清理临时文件
        for temp_img in temp_images:
            try:
                os.remove(temp_img)
            except:
                pass

    def add_pdf_screenshots_compact(self, pdf_paths: List[str], per_page: int = 2):
        """将 PDF 转为图片后紧凑添加到文档（适用于横版发票，两张一页，适当缩小）"""
        if not HAS_PYMUPDF:
            self.doc.add_paragraph(f"⚠️ 无法添加 PDF (PyMuPDF 未安装)")
            return

        temp_images = []

        for pdf_path in pdf_paths:
            try:
                doc = fitz.open(pdf_path)
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    # 渲染为图片（1.5倍清晰度，比2倍小一点以便缩小显示）
                    mat = fitz.Matrix(1.5, 1.5)
                    pix = page.get_pixmap(matrix=mat)

                    temp_img_path = f"/tmp/pdf_compact_{Path(pdf_path).stem}_{page_num}.png"
                    pix.save(temp_img_path)
                    temp_images.append(temp_img_path)
                doc.close()
            except Exception as e:
                print(f"⚠️  PDF 转换失败 {pdf_path}: {e}")

        # 每 per_page 张图片为一组，上下排列在一页
        for i in range(0, len(temp_images), per_page):
            batch = temp_images[i:i + per_page]

            for img_path in batch:
                paragraph = self.doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                try:
                    # 缩小宽度以适应两张发票在一页
                    run.add_picture(img_path, width=Cm(15))
                except Exception as e:
                    paragraph.add_run(f"[图片加载失败: {Path(img_path).name}]")

            # 每组后添加分页（除非是最后一组）
            if i + per_page < len(temp_images):
                self.add_page_break()

        # 清理临时文件
        for temp_img in temp_images:
            try:
                os.remove(temp_img)
            except:
                pass
    
    def add_page_break(self):
        """添加分页符"""
        self.doc.add_page_break()
    
    def save(self, output_path: str) -> str:
        """保存文档"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        return str(output_path)


if __name__ == "__main__":
    # 测试
    gen = WordGenerator()
    gen.add_title("加班报销材料汇总")
    gen.add_section_title("打卡记录截图")
    gen.doc.add_paragraph("（这里放置打卡截图）")
    output = gen.save("test_summary.docx")
    print(f"生成 Word 文档: {output}")
